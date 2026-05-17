"""Word document parser plugin using python-docx.

Extracts text, headings, tables, and images from .docx files.
"""

import logging
import os
import uuid

from app.services.parsers.base import Asset, Block, ParsedDocument, ParseError

logger = logging.getLogger(__name__)


class DocxParser:
    """Word document parser using python-docx.

    Extracts:
    - Paragraphs with style information (heading levels, bold, italic)
    - Tables with row/column structure
    - Inline images
    """

    name: str = "docx-parser"
    supported_extensions: list[str] = ["docx"]
    priority: int = 100

    def can_parse(self, file_path: str, mime_type: str) -> bool:
        """Check if this parser can handle the file."""
        ext = os.path.splitext(file_path)[1].lower().lstrip(".")
        return ext in self.supported_extensions or mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    async def parse(self, file_path: str) -> ParsedDocument:
        """Parse a Word document.

        Args:
            file_path: Path to the .docx file

        Returns:
            ParsedDocument with extracted content

        Raises:
            ParseError: If the file is corrupted or password-protected
        """
        if not os.path.exists(file_path):
            raise ParseError(f"File not found: {file_path}", reason="corrupted")

        try:
            from docx import Document as DocxDocument
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            raise ParseError(
                "python-docx is not installed",
                reason="unknown",
            )

        try:
            doc = DocxDocument(file_path)
        except PackageNotFoundError:
            raise ParseError(
                f"File is not a valid DOCX document: {file_path}",
                reason="corrupted",
            )
        except Exception as e:
            error_msg = str(e).lower()
            if "password" in error_msg or "encrypted" in error_msg:
                raise ParseError(
                    f"Document is password-protected: {file_path}",
                    reason="password_protected",
                )
            if "corrupt" in error_msg or "invalid" in error_msg:
                raise ParseError(
                    f"Document is corrupted: {file_path}",
                    reason="corrupted",
                )
            raise ParseError(
                f"Failed to parse DOCX: {file_path}: {e}",
                reason="unknown",
            )

        blocks: list[Block] = []
        assets: list[Asset] = []
        metadata: dict = {
            "source": file_path,
            "page_count": 1,  # DOCX doesn't have explicit pages
        }

        # Extract core properties
        if doc.core_properties:
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            block_type = "paragraph"
            style: dict = {}

            # Detect heading style
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if "heading" in style_name:
                    block_type = "heading"
                    # Extract heading level from style name (e.g., "Heading 1" -> 1)
                    try:
                        level = int(style_name.replace("heading", "").strip())
                        style["heading_level"] = min(level, 6)
                    except ValueError:
                        style["heading_level"] = 1

            # Check for bold/italic in runs
            has_bold = any(run.bold for run in para.runs if run.bold is not None)
            has_italic = any(run.italic for run in para.runs if run.italic is not None)
            style["bold"] = has_bold
            style["italic"] = has_italic

            # Check font size
            for run in para.runs:
                if run.font and run.font.size:
                    style["font_size"] = run.font.size.pt
                    break

            blocks.append(Block(
                type=block_type,
                text=text,
                page_number=1,
                style=style,
                raw={"style_name": para.style.name if para.style else ""},
            ))

        # Process tables
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            if table_text:
                blocks.append(Block(
                    type="table",
                    text=table_text,
                    page_number=1,
                    style={"table_index": table_idx},
                    raw={"rows": len(table.rows), "cols": len(table.columns)},
                ))

        # Extract images from document relationships
        assets = self._extract_images(doc)
        for asset in assets:
            blocks.append(Block(
                type="image",
                text=f"[Image: {asset.id}]",
                page_number=1,
                style={},
                raw={"asset_id": asset.id},
            ))

        if not blocks:
            raise ParseError(
                f"No content extracted from DOCX: {file_path}",
                reason="empty",
            )

        return ParsedDocument(blocks=blocks, metadata=metadata, assets=assets)

    def _extract_table_text(self, table) -> str:
        """Convert a docx table to Markdown table format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if not rows:
            return ""

        # Add header separator after first row
        if len(rows) >= 1:
            col_count = len(table.columns)
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            rows.insert(1, separator)

        return "\n".join(rows)

    def _extract_images(self, doc) -> list[Asset]:
        """Extract embedded images from the document."""
        assets: list[Asset] = []
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    try:
                        image_data = rel.target_part.blob
                        content_type = rel.target_part.content_type
                        asset_id = str(uuid.uuid4())
                        assets.append(Asset(
                            id=asset_id,
                            type="image",
                            data=image_data,
                            mime_type=content_type or "image/png",
                            page_number=1,
                        ))
                    except Exception:
                        continue
        except Exception:
            pass
        return assets
